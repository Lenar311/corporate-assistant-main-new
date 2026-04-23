# document_processor.py — обработка и индексация документов с поддержкой кодировок
# 🔥 ВЕРСИЯ: 2.0 — исправлена работа с Path на Windows

import hashlib
import json
import os
import re
from pathlib import Path
from typing import Dict, List, Optional

import chromadb
import ollama
from sentence_transformers import SentenceTransformer

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

from config import (
    CHROMA_DB_PATH,
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    DATA_DIR,
    DOCUMENTS_DIR,
    EMBEDDING_BATCH_SIZE,
    EMBEDDING_MODEL,
    INDEX_STATE_PATH,
    OLLAMA_BASE_URL,
    VISION_MODEL,
    logger,
)
from metadata_extractor import extract_metadata
from text_splitter import simple_text_splitter


def _file_hash(filepath: str) -> str:
    """SHA-256 хэш файла для отслеживания изменений."""
    # 🔥 Преобразуем Path в строку если нужно
    filepath = str(filepath) if isinstance(filepath, Path) else filepath
    h = hashlib.sha256()
    with open(filepath, "rb") as f:
        for block in iter(lambda: f.read(8192), b""):
            h.update(block)
    return h.hexdigest()


def _load_index_state() -> Dict[str, str]:
    """Загружает состояние индекса: {filename: hash}."""
    if INDEX_STATE_PATH.exists():
        with open(INDEX_STATE_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def _save_index_state(state: Dict[str, str]) -> None:
    """Сохраняет состояние индекса."""
    INDEX_STATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(INDEX_STATE_PATH, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)


def _extract_text_from_docx(filepath: str) -> str:
    """Извлекает текст из .docx файла с помощью python-docx."""
    # 🔥 Преобразуем Path в строку если нужно
    filepath = str(filepath) if isinstance(filepath, Path) else filepath
    
    if not DOCX_AVAILABLE:
        logger.error("Установите python-docx: pip install python-docx")
        return ""
    try:
        doc = DocxDocument(filepath)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))
        return "\n".join(parts).strip()
    except Exception as e:
        logger.error("Ошибка при чтении DOCX %s: %s", filepath, e)
        return ""


def _extract_text_from_doc(filepath: str) -> str:
    """Извлекает текст из .doc файла через Word COM (Windows)."""
    # 🔥 Преобразуем Path в строку если нужно
    filepath = str(filepath) if isinstance(filepath, Path) else filepath
    filepath = os.path.abspath(filepath)
    
    try:
        import win32com.client  # type: ignore
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(filepath, ReadOnly=True, AddToRecentFiles=False)
        text = doc.Content.Text
        doc.Close(SaveChanges=False)
        word.Quit()
        del doc, word
        return text.strip() if text and len(text.strip()) > 100 else ""
    except ImportError:
        logger.error("pywin32 не установлен: pip install pywin32")
        return ""
    except Exception as e:
        logger.error("COM-ошибка при обработке %s: %s", filepath, e)
        try:
            import psutil
            for proc in psutil.process_iter(['name']):
                if proc.info['name'] and 'WINWORD.EXE' in proc.info['name'].upper():
                    proc.kill()
        except:
            pass
        return ""


def _extract_text_from_txt(filepath: str) -> str:
    """Извлекает текст из TXT-файла с автоопределением кодировки."""
    # 🔥 Преобразуем Path в строку если нужно
    filepath = str(filepath) if isinstance(filepath, Path) else filepath
    
    for encoding in ['utf-8-sig', 'cp1251', 'utf-8', 'latin-1', 'cp866']:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                text = f.read().strip()
                if text:
                    logger.info("✅ Прочитан TXT %s в кодировке %s", os.path.basename(filepath), encoding)
                    return text
        except Exception:
            continue
    logger.error("❌ Не удалось определить кодировку TXT %s", filepath)
    return ""


# def _extract_text_from_rtf(filepath: str) -> str:
#     """Извлекает текст из RTF-файла через striprtf."""
#     # 🔥 Преобразуем Path в строку если нужно
#     filepath = str(filepath) if isinstance(filepath, Path) else filepath
    
#     try:
#         from striprtf.striprtf import rtf_to_text
#         with open(filepath, 'rb') as f:
#             content = f.read()
#         for encoding in ['utf-8', 'cp1251', 'latin-1']:
#             try:
#                 decoded = content.decode(encoding)
#                 text = rtf_to_text(decoded).strip()
#                 if text:
#                     logger.info("✅ Прочитан RTF %s в кодировке %s", os.path.basename(filepath), encoding)
#                     return text
#             except:
#                 continue
#         return ""
#     except ImportError:
#         logger.error("Установите striprtf: pip install striprtf")
#         return ""
#     except Exception as e:
#         logger.error("Ошибка при чтении RTF %s: %s", filepath, e)
#         return ""


class DocumentProcessor:
    def __init__(self):
        logger.info("📦 Загрузка эмбеддинг-модели: %s", EMBEDDING_MODEL)
        self.embedding_model = SentenceTransformer(EMBEDDING_MODEL)
        logger.info("✅ Эмбеддинг-модель загружена.")
        
        self.ollama_client = ollama.Client(host=OLLAMA_BASE_URL)
        
        # 🔥 Преобразуем Path в строку для ChromaDB
        chroma_path_str = str(CHROMA_DB_PATH)
        os.makedirs(chroma_path_str, exist_ok=True)
        self.client = chromadb.PersistentClient(path=chroma_path_str)
        self.collection = self.client.get_or_create_collection(name="legal_docs")

    def vision_extract_pdf(self, filepath: str) -> str:
        """Извлекает текст из PDF с помощью vision-модели (для сканов)."""
        # 🔥 Преобразуем Path в строку если нужно
        filepath = str(filepath) if isinstance(filepath, Path) else filepath
        
        if not PYMUPDF_AVAILABLE:
            logger.error("Установите PyMuPDF: pip install PyMuPDF")
            return ""
        
        vision_prompt = (
            "Извлеки весь текст с этой страницы документа. "
            "Если есть таблицы — воспроизведи их содержимое в текстовом виде. "
            "Отвечай только на русском языке."
        )
        try:
            doc = fitz.open(filepath)
            full_text = ""
            for i, page in enumerate(doc):
                logger.info("  Vision-обработка страницы %d/%d", i + 1, len(doc))
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                response = self.ollama_client.chat(
                    model=VISION_MODEL,
                    messages=[{"role": "user", "content": vision_prompt, "images": [img_bytes]}],
                )
                full_text += f"\n--- Страница {i + 1} ---\n{response.message.content}"
            doc.close()
            return full_text
        except Exception as e:
            logger.error("Vision ошибка для %s: %s", filepath, e)
            return ""

    @staticmethod
    def _is_readable_text(text: str) -> bool:
        """Проверяет, что текст читаемый."""
        if not text or len(text.strip()) < 50:
            return False
        alpha_count = sum(1 for c in text if c.isalpha() or c.isdigit() or c.isspace())
        return alpha_count / len(text) > 0.5

    def process_pdf(self, filepath: str) -> List[Dict]:
        """Обрабатывает PDF-файл."""
        # 🔥 Преобразуем Path в строку если нужно
        filepath = str(filepath) if isinstance(filepath, Path) else filepath
        
        text = ""
        if PDFMINER_AVAILABLE:
            try:
                text = pdfminer_extract_text(filepath) or ""
            except Exception as e:
                logger.info("pdfminer не смог извлечь текст из %s: %s", filepath, type(e).__name__)
        
        if not self._is_readable_text(text):
            logger.info("Текст нечитаемый, переключение на vision-модель для %s", filepath)
            text = self.vision_extract_pdf(filepath)
        
        if not text.strip():
            raise ValueError(f"Текст не извлечён из {filepath}")
        
        metadata = extract_metadata(text, os.path.basename(filepath))
        return self._semantic_chunking(text, metadata, filepath)

    def process_doc(self, filepath: str) -> List[Dict]:
        """Обрабатывает DOC/DOCX-файл."""
        # 🔥 Преобразуем Path в строку если нужно
        filepath = str(filepath) if isinstance(filepath, Path) else filepath
        
        ext = Path(filepath).suffix.lower()
        if ext == ".docx":
            text = _extract_text_from_docx(filepath)
        elif ext == ".doc":
            text = _extract_text_from_doc(filepath)
        else:
            raise ValueError(f"Неподдерживаемый формат: {ext}")
        
        if not text or not text.strip():
            raise ValueError(f"Не удалось извлечь текст из {filepath}")
        
        metadata = extract_metadata(text, os.path.basename(filepath))
        return self._semantic_chunking(text, metadata, filepath)

    def _semantic_chunking(self, text: str, metadata: Dict, filepath: str) -> List[Dict]:
        """Разбивает текст на семантические чанки."""
        if not text or not text.strip():
            logger.warning("Пустой текст для %s", filepath)
            return []

        section_pattern = r"\n+(?=\s*(?:\d+(?:\.\d+)*|Глава\s+\d+|Раздел\s+[IVXLCDM]+|Пункт\s+\d+|Статья\s+\d+)\b)"
        sections = re.split(section_pattern, text)
        chunks: List[Dict] = []
        chunk_id = 0

        for section in sections:
            section = section.strip()
            if "приложение" in section.lower()[:50]:
                logger.info("Найдено приложение в разделе: %s", section[:100])
                continue

            section_title = ""
            lines = [line.strip() for line in section.split('\n') if line.strip()]
            if lines:
                first_line = lines[0]
                first_line = re.sub(r'[\xa0\u200b]+', ' ', first_line).strip()
                if len(first_line) <= 120 and any(kw in first_line.lower() for kw in [
                    'ведомость', 'спецификация', 'требования', 'оформление', 'шрифт',
                    'нормоконтроль', 'дублирование', 'покупные изделия', 'графы'
                ]):
                    section_title = first_line

            chunk_texts = simple_text_splitter(section, CHUNK_SIZE, CHUNK_OVERLAP)
            for chunk_text in chunk_texts:
                clean_text = re.sub(r"\s+", " ", chunk_text).strip()
                if len(clean_text) < 50:
                    continue

                clause_number = None
                first_line = clean_text.split('\n', 1)[0].strip()
                clause_match = re.match(r'^(\d{1,2}(?:\.\d{1,2}){1,3})', first_line)
                if clause_match:
                    clause_number = clause_match.group(1)

                start_pos = text.find(chunk_text)
                end_pos = start_pos + len(chunk_text) if start_pos != -1 else -1

                chunk_meta = metadata.copy()
                chunk_meta.update({
                    "section_title": section_title,
                    "clause": clause_number,
                    "section_id": f"sec_{chunk_id}",
                    "char_start": start_pos if start_pos != -1 else -1,
                    "char_end": end_pos if end_pos != -1 else -1,
                    "filepath": os.path.abspath(filepath),
                    "chunk_index": chunk_id,
                })

                chunks.append({
                    "text": clean_text,
                    "metadata": chunk_meta,
                    "id": f"{metadata['filename']}_{chunk_id}",
                })
                chunk_id += 1

        return chunks

    def _remove_file_from_index(self, filename: str) -> None:
        """Удаляет из ChromaDB все записи, связанные с файлом."""
        try:
            results = self.collection.get(where={"filename": filename})
            if results["ids"]:
                self.collection.delete(ids=results["ids"])
                logger.info("Удалено %d записей для %s", len(results["ids"]), filename)
        except Exception as e:
            logger.warning("Не удалось удалить записи для %s: %s", filename, e)

    def _cleanup_orphan_chunks(self, current_filenames: set) -> None:
        """Удаляет устаревшие чанки."""
        try:
            all_records = self.collection.get()
            if not all_records["ids"]:
                return
            orphan_ids = []
            for rec_id, meta in zip(all_records["ids"], all_records["metadatas"]):
                if meta and meta.get("filename") not in current_filenames:
                    orphan_ids.append(rec_id)
            if orphan_ids:
                self.collection.delete(ids=orphan_ids)
                logger.info("Удалено %d устаревших записей", len(orphan_ids))
        except Exception as e:
            logger.warning("Ошибка при очистке: %s", e)

    def index_documents(self) -> Dict[str, int]:
        """Индексирует документы из DOCUMENTS_DIR."""
        stats = {"processed": 0, "skipped": 0, "chunks": 0}
        
        # 🔥 ПРЕОБРАЗУЕМ PATH В СТРОКУ
        docs_dir = str(DOCUMENTS_DIR)
        
        if not os.path.exists(docs_dir):
            os.makedirs(docs_dir, exist_ok=True)
            logger.info("📁 Папка %s создана.", docs_dir)
            return stats

        files = [f for f in os.listdir(docs_dir) if f.lower().endswith((".pdf", ".doc", ".docx", ".txt"))]
        if not files:
            logger.info("📁 Папка %s пуста.", docs_dir)
            return stats

        self._cleanup_orphan_chunks(set(files))

        index_state = _load_index_state()
        current_files: Dict[str, str] = {}
        files_to_process: List[str] = []

        for file in files:
            filepath = os.path.join(docs_dir, file)
            file_h = _file_hash(filepath)
            current_files[file] = file_h

            if index_state.get(file) == file_h:
                logger.info("✓ Пропуск: %s", file)
                stats["skipped"] += 1
                continue
            files_to_process.append(file)

        if not files_to_process:
            logger.info("✅ Все документы уже проиндексированы.")
            _save_index_state(current_files)
            return stats

        all_chunks: List[Dict] = []
        for file in files_to_process:
            filepath = os.path.join(docs_dir, file)
            logger.info("📄 Обработка: %s", file)
            self._remove_file_from_index(file)
            ext = Path(filepath).suffix.lower()
            try:
                if ext == ".pdf":
                    chunks = self.process_pdf(filepath)
                elif ext in (".doc", ".docx"):
                    chunks = self.process_doc(filepath)
                elif ext == ".txt":
                    text = _extract_text_from_txt(filepath)
                    if not text.strip():
                        raise ValueError("Пустой TXT-файл")
                    metadata = extract_metadata(text, os.path.basename(filepath))
                    chunks = self._semantic_chunking(text, metadata, filepath)
                else:
                    continue
                
                all_chunks.extend(chunks)
                stats["processed"] += 1
                stats["chunks"] += len(chunks)
                
            except Exception as e:
                logger.error("❌ Ошибка при обработке %s: %s", file, e)
                continue

        if not all_chunks:
            logger.warning("⚠️ Нет данных для индексации.")
            _save_index_state(current_files)
            return stats

        logger.info("🔢 Генерация эмбеддингов для %d фрагментов...", len(all_chunks))
        
        texts = [chunk["text"] for chunk in all_chunks]
        
        # 🔥 Для BGE-M3 НЕ нужен префикс "passage: "
        # Проверяем тип модели
        if "e5" in EMBEDDING_MODEL.lower():
            prefixed_texts = [f"passage: {text}" for text in texts]
        else:
            # BGE-M3 и другие модели без префикса
            prefixed_texts = texts

        embeddings = self.embedding_model.encode(
            prefixed_texts,
            batch_size=EMBEDDING_BATCH_SIZE,
            show_progress_bar=True,
            convert_to_numpy=True,
            normalize_embeddings=True
        ).tolist()

        metadatas = []
        for chunk in all_chunks:
            meta = chunk["metadata"].copy()
            for k, v in meta.items():
                if v is None:
                    meta[k] = ""
            metadatas.append(meta)
        
        ids = [chunk["id"] for chunk in all_chunks]

        self.collection.add(
            documents=texts,
            embeddings=embeddings,
            metadatas=metadatas,
            ids=ids,
        )

        _save_index_state(current_files)
        logger.info("✅ Проиндексировано %d фрагментов из %d документов.", len(all_chunks), len(files_to_process))
        return stats